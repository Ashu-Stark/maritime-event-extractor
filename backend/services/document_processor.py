"""
Document Processing Service
Handles text extraction from PDF, DOC, and DOCX files
"""

import os
import logging
from typing import Dict, Any
try:
    import PyPDF2  # type: ignore
except Exception:
    PyPDF2 = None  # type: ignore
try:
    from docx import Document as DocxDocument  # type: ignore
except Exception:
    DocxDocument = None  # type: ignore
import tempfile
import subprocess
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing and extracting text from various document formats"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf', '.doc', '.docx', '.txt'
        }
        
    def extract_text(self, file_path: str) -> str:
        """
        Extract text content from document
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            Exception: If extraction fails
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(
                    f"Unsupported file format: {file_extension}"
                )
            
            logger.info(f"Extracting text from {file_extension} file: {file_path}")
            
            if file_extension == '.pdf':
                if PyPDF2 is None:
                    raise ImportError("PyPDF2 is required to process PDF files")
                return self._extract_from_pdf(file_path)
            elif file_extension == '.docx':
                if DocxDocument is None:
                    raise ImportError("python-docx is required to process DOCX files")
                return self._extract_from_docx(file_path)
            elif file_extension == '.doc':
                return self._extract_from_doc(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
            
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using multiple methods"""
        try:
            # Method 1: Try PyPDF2 first
            if PyPDF2 is not None:
                try:
                    text_content = []
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        
                        for page_num, page in enumerate(pdf_reader.pages):
                            try:
                                page_text = page.extract_text()
                                if page_text.strip():
                                    text_content.append(
                                        f"--- Page {page_num + 1} ---\n{page_text}"
                                    )
                            except Exception as e:
                                logger.warning(
                                    f"Failed to extract text from page "
                                    f"{page_num + 1}: {str(e)}"
                                )
                                continue
                    
                    if text_content:
                        logger.info("PyPDF2 extraction successful")
                        return '\n\n'.join(text_content)
                except Exception as e:
                    logger.warning(f"PyPDF2 extraction failed: {str(e)}")
            
            # Method 2: Try PyMuPDF (fitz)
            try:
                import fitz
                text_content = []
                doc = fitz.open(file_path)
                
                for page_num in range(len(doc)):
                    try:
                        page = doc.load_page(page_num)
                        page_text = page.get_text()
                        if page_text.strip():
                            text_content.append(
                                f"--- Page {page_num + 1} ---\n{page_text}"
                            )
                    except Exception as e:
                        logger.warning(
                            f"PyMuPDF failed for page {page_num + 1}: {str(e)}"
                        )
                        continue
                
                doc.close()
                
                if text_content:
                    logger.info("PyMuPDF extraction successful")
                    return '\n\n'.join(text_content)
            except ImportError:
                logger.info("PyMuPDF not available")
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed: {str(e)}")
            
            # Method 3: Try pdfplumber
            try:
                import pdfplumber
                text_content = []
                
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                text_content.append(
                                    f"--- Page {page_num + 1} ---\n{page_text}"
                                )
                        except Exception as e:
                            logger.warning(
                                f"pdfplumber failed for page {page_num + 1}: {str(e)}"
                            )
                            continue
                
                if text_content:
                    logger.info("pdfplumber extraction successful")
                    return '\n\n'.join(text_content)
            except ImportError:
                logger.info("pdfplumber not available")
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {str(e)}")
            
            # Method 4: Try OCR as final fallback
            logger.warning("All direct text extraction methods failed, trying OCR")
            return self._try_ocr_extraction(file_path)
            
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            if DocxDocument is None:
                raise ImportError("python-docx is not installed")
            doc = DocxDocument(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(' | '.join(row_text))
                
                if table_text:
                    text_content.append('\n--- Table ---\n' + '\n'.join(table_text))
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise
    
    def _extract_from_doc(self, file_path: str) -> str:
        """Extract text from DOC file using antiword or similar tool"""
        try:
            # Try using antiword (Linux/Mac) or textract (cross-platform)
            result = None
            
            # Method 1: Try antiword (if available)
            try:
                result = subprocess.run(
                    ['antiword', file_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout
            except (subprocess.TimeoutExpired, FileNotFoundError):
                pass
            
            # Method 2: Try catdoc (if available)
            try:
                result = subprocess.run(
                    ['catdoc', file_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return result.stdout
            except (subprocess.TimeoutExpired, FileNotFoundError):
                pass
            
            # Method 3: Try python-docx2txt (fallback)
            try:
                import docx2txt
                text = docx2txt.process(file_path)
                if text and text.strip():
                    return text
            except ImportError:
                pass
            
            # Method 4: Convert to DOCX and extract (using LibreOffice if available)
            try:
                return self._convert_doc_to_docx_and_extract(file_path)
            except:
                pass
            
            raise Exception("No suitable DOC extraction method available")
            
        except Exception as e:
            logger.error(f"DOC extraction error: {str(e)}")
            raise
    
    def _convert_doc_to_docx_and_extract(self, file_path: str) -> str:
        """Convert DOC to DOCX using LibreOffice and extract text"""
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                # Convert DOC to DOCX using LibreOffice
                result = subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'docx',
                    '--outdir', temp_dir, file_path
                ], capture_output=True, timeout=60)
                
                if result.returncode != 0:
                    raise Exception("LibreOffice conversion failed")
                
                # Find the converted DOCX file
                doc_name = Path(file_path).stem
                docx_path = os.path.join(temp_dir, f"{doc_name}.docx")
                
                if os.path.exists(docx_path):
                    return self._extract_from_docx(docx_path)
                else:
                    raise Exception("Converted DOCX file not found")
                    
        except Exception as e:
            logger.error(f"DOC to DOCX conversion error: {str(e)}")
            raise
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from a plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return text
        except Exception as e:
            logger.error(f"TXT extraction error: {str(e)}")
            raise
    
    def _try_ocr_extraction(self, file_path: str) -> str:
        """Try OCR extraction for image-based PDFs"""
        try:
            import pytesseract
            from pdf2image import convert_from_path
            from PIL import Image
            
            logger.info("Attempting OCR extraction for image-based PDF")
            
            # Check if Tesseract is available
            try:
                tesseract_version = pytesseract.get_tesseract_version()
                logger.info(f"Tesseract version: {tesseract_version}")
            except Exception as e:
                logger.error(f"Tesseract not available: {str(e)}")
                return "OCR_UNAVAILABLE: Tesseract OCR engine is not installed. Please install Tesseract for Windows to enable OCR processing of image-based PDFs. Download from: https://github.com/UB-Mannheim/tesseract/wiki"
            
            # Convert PDF pages to images
            try:
                images = convert_from_path(file_path, dpi=300)
                logger.info(f"Converted PDF to {len(images)} images for OCR")
            except Exception as e:
                logger.warning(f"PDF to image conversion failed: {str(e)}")
                return "OCR_FAILED: Unable to convert PDF pages to images for OCR processing. This may be due to missing Poppler utilities."
            
            # Extract text from each image using OCR
            extracted_texts = []
            for i, image in enumerate(images):
                try:
                    # Use pytesseract for OCR
                    text = pytesseract.image_to_string(image, lang='eng')
                    if text.strip():
                        extracted_texts.append(f"--- Page {i+1} ---\n{text.strip()}")
                        logger.info(f"OCR successful for page {i+1}")
                    else:
                        logger.warning(f"OCR returned empty text for page {i+1}")
                except Exception as e:
                    logger.warning(f"OCR failed for page {i+1}: {str(e)}")
                    continue
            
            if extracted_texts:
                full_text = '\n\n'.join(extracted_texts)
                logger.info(f"OCR extraction successful: {len(full_text)} characters extracted")
                return full_text
            else:
                logger.warning("OCR extraction failed for all pages")
                return "OCR_FAILED: Text extraction failed for all document pages."
                
        except ImportError as e:
            logger.error(f"OCR libraries not available: {str(e)}")
            return "OCR_UNAVAILABLE: OCR processing requires pytesseract, pdf2image, and Pillow libraries. Please install them for image-based PDF processing."
        except Exception as e:
            logger.error(f"OCR extraction error: {str(e)}")
            return f"OCR_ERROR: OCR processing failed with error: {str(e)}"
    
    def get_document_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from document
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing document metadata
        """
        try:
            file_extension = Path(file_path).suffix.lower()
            metadata = {
                'file_size': os.path.getsize(file_path),
                'file_extension': file_extension,
                'pages': 0,
                'word_count': 0,
                'character_count': 0
            }
            
            if file_extension == '.pdf':
                metadata.update(self._get_pdf_metadata(file_path))
            elif file_extension == '.docx':
                metadata.update(self._get_docx_metadata(file_path))
            
            # Extract text to get word/character counts
            try:
                text = self.extract_text(file_path)
                metadata['word_count'] = len(text.split())
                metadata['character_count'] = len(text)
            except:
                pass
            
            return metadata
            
        except Exception as e:
            logger.error(f"Metadata extraction error: {str(e)}")
            return {'error': str(e)}
    
    def _get_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract PDF-specific metadata"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata = {
                    'pages': len(pdf_reader.pages),
                    'title': '',
                    'author': '',
                    'subject': '',
                    'creator': ''
                }
                
                # Extract document info if available
                if pdf_reader.metadata:
                    info = pdf_reader.metadata
                    metadata.update({
                        'title': info.get('/Title', ''),
                        'author': info.get('/Author', ''),
                        'subject': info.get('/Subject', ''),
                        'creator': info.get('/Creator', ''),
                        'creation_date': info.get('/CreationDate', ''),
                        'modification_date': info.get('/ModDate', '')
                    })
                
                return metadata
                
        except Exception as e:
            logger.error(f"PDF metadata extraction error: {str(e)}")
            return {}
    
    def _get_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract DOCX-specific metadata"""
        try:
            doc = DocxDocument(file_path)
            
            metadata = {
                'pages': 1,  # DOCX doesn't have fixed pages
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            
            # Extract core properties if available
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                metadata.update({
                    'title': props.title or '',
                    'author': props.author or '',
                    'subject': props.subject or '',
                    'created': props.created.isoformat() if props.created else '',
                    'modified': props.modified.isoformat() if props.modified else ''
                })
            
            return metadata
            
        except Exception as e:
            logger.error(f"DOCX metadata extraction error: {str(e)}")
            return {}
    
    def validate_document(self, file_path: str) -> Dict[str, Any]:
        """
        Validate document and check if it's processable
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary with validation results
        """
        try:
            validation_result = {
                'is_valid': False,
                'file_exists': False,
                'is_supported_format': False,
                'is_readable': False,
                'has_text_content': False,
                'file_size': 0,
                'error_message': None
            }
            
            # Check if file exists
            if not os.path.exists(file_path):
                validation_result['error_message'] = "File does not exist"
                return validation_result
            
            validation_result['file_exists'] = True
            validation_result['file_size'] = os.path.getsize(file_path)
            
            # Check file format
            file_extension = Path(file_path).suffix.lower()
            if file_extension not in self.supported_formats:
                validation_result['error_message'] = f"Unsupported format: {file_extension}"
                return validation_result
            
            validation_result['is_supported_format'] = True
            
            # Try to read the file
            try:
                with open(file_path, 'rb') as f:
                    f.read(1024)  # Try to read first 1KB
                validation_result['is_readable'] = True
            except Exception as e:
                validation_result['error_message'] = f"File is not readable: {str(e)}"
                return validation_result
            
            # Try to extract some text
            try:
                text = self.extract_text(file_path)
                if text and len(text.strip()) > 10:  # At least 10 characters
                    validation_result['has_text_content'] = True
                    validation_result['is_valid'] = True
                else:
                    validation_result['error_message'] = "No readable text content found"
            except Exception as e:
                validation_result['error_message'] = f"Text extraction failed: {str(e)}"
            
            return validation_result
            
        except Exception as e:
            logger.error(f"Document validation error: {str(e)}")
            return {
                'is_valid': False,
                'error_message': f"Validation failed: {str(e)}"
            }
